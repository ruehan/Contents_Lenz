import os
import PyPDF2
import docx
from pathlib import Path
from src.config import SUPPORTED_TEXT_FORMATS, SUPPORTED_DOCUMENT_FORMATS

class FileHandler:
    """파일 처리를 위한 유틸리티 클래스"""
    
    @staticmethod
    def read_file(file_path):
        """파일 경로에서 텍스트 내용을 읽어옵니다.
        
        Args:
            file_path (str): 읽을 파일의 경로
            
        Returns:
            str: 파일의 텍스트 내용
            
        Raises:
            ValueError: 지원하지 않는 파일 형식일 경우
            FileNotFoundError: 파일을 찾을 수 없을 경우
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        # 텍스트 파일 처리
        if file_extension in SUPPORTED_TEXT_FORMATS:
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except UnicodeDecodeError:
                # UTF-8로 읽기 실패 시 다른 인코딩 시도
                try:
                    with open(file_path, 'r', encoding='cp949') as file:
                        return file.read()
                except UnicodeDecodeError:
                    raise ValueError(f"파일 인코딩을 인식할 수 없습니다: {file_path}")
        
        # PDF 파일 처리
        elif file_extension == '.pdf':
            return FileHandler._read_pdf(file_path)
        
        # Word 문서 처리
        elif file_extension in ['.docx', '.doc']:
            return FileHandler._read_docx(file_path)
        
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_extension}")
    
    @staticmethod
    def _read_pdf(file_path):
        """PDF 파일에서 텍스트를 추출합니다.
        
        Args:
            file_path (Path): PDF 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            raise ValueError(f"PDF 파일을 읽는 중 오류가 발생했습니다: {str(e)}")
    
    @staticmethod
    def _read_docx(file_path):
        """Word 문서에서 텍스트를 추출합니다.
        
        Args:
            file_path (Path): Word 문서 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            doc = docx.Document(file_path)
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            return text
        except Exception as e:
            raise ValueError(f"Word 문서를 읽는 중 오류가 발생했습니다: {str(e)}")
    
    @staticmethod
    def save_text(text, file_path, file_format='.txt'):
        """텍스트를 파일로 저장합니다.
        
        Args:
            text (str): 저장할 텍스트
            file_path (str): 저장할 파일 경로 (확장자 없이)
            file_format (str): 파일 형식 ('.txt', '.pdf', '.docx')
            
        Returns:
            str: 저장된 파일의 전체 경로
            
        Raises:
            ValueError: 지원하지 않는 출력 형식일 경우
        """
        file_path = Path(file_path)
        full_path = file_path.with_suffix(file_format)
        
        # 디렉토리가 없으면 생성
        os.makedirs(file_path.parent, exist_ok=True)
        
        # 텍스트 파일로 저장
        if file_format == '.txt':
            with open(full_path, 'w', encoding='utf-8') as file:
                file.write(text)
            return str(full_path)
        
        # Word 문서로 저장
        elif file_format == '.docx':
            doc = docx.Document()
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            doc.save(full_path)
            return str(full_path)
        
        # PDF로 저장 (추후 구현)
        elif file_format == '.pdf':
            # PDF 생성 라이브러리 추가 필요 (예: reportlab)
            raise NotImplementedError("PDF 저장 기능은 아직 구현되지 않았습니다.")
        
        else:
            raise ValueError(f"지원하지 않는 출력 형식입니다: {file_format}") 